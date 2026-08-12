from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\cwbec\BlockMed")
MARKDOWN = ROOT / "docs" / "Blockmediary_Deal_Value_Research_Report_Submission_Edition.md"
OUTPUT = ROOT / "docs" / "Blockmediary_Deal_Value_Research_Report_Submission_Edition.docx"
ASSETS = ROOT / "tmp" / "pdfs" / "deal_value_assets"

NAVY = "193652"
TEAL = "1B7F7A"
INK = "253142"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_TEAL = "EAF5F4"
WHITE = "FFFFFF"
GOLD = "B88120"


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(element, fill):
    tc_pr = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def current_content_width_dxa(section):
    return int(section.page_width - section.left_margin - section.right_margin)


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text, url, bold=False, italic=False, color=TEAL):
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\(https?://[^)]+\)|\[S\d+\])")


def add_inline(paragraph, text, base_size=11, base_color=INK, italic_all=False):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size=base_size, color=base_color, italic=italic_all)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=base_size, bold=True, color=base_color, italic=italic_all)
        elif token.startswith("[") and "](" in token:
            label, url = token[1:].split("](", 1)
            add_hyperlink(paragraph, label, url[:-1], italic=italic_all)
        else:
            run = paragraph.add_run(token)
            set_font(run, size=base_size, bold=True, color=TEAL, italic=italic_all)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=base_size, color=base_color, italic=italic_all)


def create_numbering(doc, kind):
    root = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    root.insert(0, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    root.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def add_list_paragraph(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    apply_numbering(p, num_id)
    add_inline(p, text)
    return p


def set_paragraph_border(paragraph, color=TEAL, size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    borders.append(left)


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.12
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_TEAL)
    p_pr.append(shd)
    set_paragraph_border(p)
    add_inline(p, text, base_size=10.5)
    return p


def split_md_row(line):
    line = line.strip().strip("|")
    return [cell.strip().replace("\\|", "|") for cell in line.split("|")]


def column_widths(rows, total):
    cols = len(rows[0])
    scores = []
    for col in range(cols):
        values = [len(re.sub(r"\[[^\]]+\]\([^)]+\)", "link", row[col])) for row in rows]
        score = max(7, min(max(values, default=7), 46))
        scores.append(score)
    minimum = 720 if cols >= 7 else 900 if cols >= 5 else 1200
    widths = [max(minimum, int(total * score / sum(scores))) for score in scores]
    difference = total - sum(widths)
    widths[-1] += difference
    if widths[-1] < minimum:
        deficit = minimum - widths[-1]
        widths[-1] = minimum
        donors = sorted(range(cols - 1), key=lambda i: widths[i], reverse=True)
        for idx in donors:
            available = max(0, widths[idx] - minimum)
            take = min(available, deficit)
            widths[idx] -= take
            deficit -= take
            if deficit == 0:
                break
    return widths


def is_numeric_cell(text):
    clean = text.replace(",", "").strip()
    return bool(re.fullmatch(r"[-+]?[$£]?[0-9.]+(?:%|k|m|bn|x)?", clean, re.I))


def add_markdown_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    total = current_content_width_dxa(doc.sections[-1])
    widths = column_widths(rows, total)
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    font_size = 7.2 if cols >= 8 else 8.0 if cols >= 7 else 8.6 if cols >= 6 else 9.2 if cols >= 4 else 10
    for ri, source_row in enumerate(rows):
        for ci, text in enumerate(source_row):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_numeric_cell(text) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, base_size=font_size, base_color=WHITE if ri == 0 else INK)
            if ri == 0:
                shade(cell._tc, NAVY)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            elif ri % 2 == 0:
                shade(cell._tc, LIGHT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = False


def add_page_furniture(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    left = p.add_run("BLOCKMEDIARY COMMERCIAL RESEARCH")
    set_font(left, size=8, bold=True, color=TEAL)
    right = p.add_run("\tSUBMISSION EDITION")
    set_font(right, size=8, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    left = p.add_run("Blockmediary commercial research | 11 August 2026")
    set_font(left, size=8.5, color=MUTED)
    page_label = p.add_run("\tPage ")
    set_font(page_label, size=8.5, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(58)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("BLOCKMEDIARY COMMERCIAL RESEARCH")
    set_font(run, size=10, bold=True, color=TEAL)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ppr = title._p.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    title.add_run("Defensible Average Deal Values for Documentary Escrow")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Year-1 base cases and Year-2 to Year-5 growth assumptions for Tier A, Tier B and Tier C document verification")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(26)
    add_inline(meta, "**Prepared by:** Transakt (Blockmediary project team)\n**Research date:** 11 August 2026\n**Primary launch market:** UAE", base_size=10, base_color=MUTED)
    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout.paragraph_format.left_indent = Inches(0.65)
    callout.paragraph_format.right_indent = Inches(0.65)
    callout.paragraph_format.space_before = Pt(28)
    callout.paragraph_format.space_after = Pt(16)
    callout.paragraph_format.line_spacing = 1.25
    p_pr = callout._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_TEAL)
    p_pr.append(shd)
    add_inline(callout, "**Recommended Year 1**\nTier A £35,000  |  Tier B £30,000  |  Tier C £30,000\nBase annual growth: 10%  |  8%  |  5%", base_size=12, base_color=NAVY)
    doc.add_page_break()


def add_contents(doc, major_headings):
    doc.add_paragraph("Contents", style="Heading 1")
    for heading in major_headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(heading)
        set_font(run, size=11, color=TEAL)
    doc.add_page_break()


def make_landscape(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_page_furniture(section)


def make_portrait(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_furniture(section)


def parse_body(doc, lines, bullet_num, decimal_num):
    i = 0
    pending_figure = None
    first_major = True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("<!-- figure:"):
            pending_figure = int(re.search(r"(\d+)", stripped).group(1))
            i += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:]
            if heading.startswith("Appendix A"):
                make_landscape(doc)
            elif heading.startswith("Appendix B"):
                make_portrait(doc)
            elif not first_major:
                doc.add_page_break()
            first_major = False
            doc.add_paragraph(heading, style="Heading 1")
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
            i += 1
            continue
        if stripped.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [split_md_row(x) for x in table_lines if not re.fullmatch(r"\|?[\s:|-]+\|?", x)]
            add_markdown_table(doc, rows)
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:])
            i += 1
            continue
        if re.match(r"^- ", stripped):
            add_list_paragraph(doc, stripped[2:], bullet_num)
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            add_list_paragraph(doc, re.sub(r"^\d+\. ", "", stripped), decimal_num)
            i += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*"):
            caption_text = stripped[1:-1]
            if pending_figure is not None:
                image_path = ASSETS / f"asset-{pending_figure:02d}.png"
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(6.15 if doc.sections[-1].orientation == WD_ORIENT.PORTRAIT else 8.7))
                pending_figure = None
            p = doc.add_paragraph(style="Caption")
            add_inline(p, caption_text, base_size=9, base_color=MUTED)
            i += 1
            continue
        if stripped.startswith("# ") or stripped.startswith("**Prepared by:**"):
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


def main():
    text = MARKDOWN.read_text(encoding="utf-8")
    lines = text.splitlines()
    if lines and lines[0].strip() == "---":
        end = lines.index("---", 1)
        lines = lines[end + 1 :]
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Executive conclusion")
    body = lines[start:]
    major_headings = [line.strip()[3:] for line in body if line.strip().startswith("## ")]

    doc = Document()
    compat = doc.settings._element.find(qn("w:compat"))
    if compat is not None:
        for setting in compat.findall(qn("w:compatSetting")):
            if setting.get(qn("w:name")) == "compatibilityMode":
                setting.set(qn("w:val"), "15")
    configure_styles(doc)
    props = doc.core_properties
    props.title = "Defensible Average Deal Values for Documentary Escrow"
    props.subject = "Blockmediary commercial research: deal-value assumptions and model implications"
    props.author = "Transakt (Blockmediary project team)"
    props.last_modified_by = "Transakt (Blockmediary project team)"
    props.keywords = "Blockmediary, documentary escrow, deal value, UAE trade, eBL, unit economics"
    props.comments = "Prepared by Transakt for academic and investor review."

    add_page_furniture(doc.sections[0])
    add_cover(doc)
    add_contents(doc, major_headings)
    bullet_num = create_numbering(doc, "bullet")
    decimal_num = create_numbering(doc, "decimal")
    parse_body(doc, body, bullet_num, decimal_num)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
