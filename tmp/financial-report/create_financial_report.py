from __future__ import annotations

import json
import math
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"C:\Users\cwbec\BlockMed")
TMP = ROOT / "tmp" / "financial-report"
DATA_PATH = TMP / "report-data.json"
ASSET_DIR = TMP / "report-assets"
OUTPUT_DOCX = ROOT / "docs" / "PDFs For Submission" / "Blockmediary_Financial_Value_and_Commercial_Viability_Report.docx"
TABLE_HELPER_DIR = Path(r"C:\Users\cwbec\.codex\plugins\cache\openai-primary-runtime\documents\26.812.11052\skills\documents\scripts")
sys.path.insert(0, str(TABLE_HELPER_DIR))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


NAVY = "19344D"
DEEP_NAVY = "10263B"
TEAL = "23847F"
GOLD = "C58E2C"
GREEN = "2F7D62"
RED = "A33A3A"
INK = "203040"
MUTED = "66758A"
LIGHT_BLUE = "EAF0F5"
LIGHT_TEAL = "E7F2F0"
LIGHT_GOLD = "FBF2DF"
LIGHT_GREEN = "E8F3ED"
LIGHT_RED = "F8EAEA"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D7DEE6"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, name="Calibri", size=None, bold=None, italic=None, color=None, all_caps=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)
    if all_caps is not None:
        run.font.all_caps = all_caps


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, spec in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key, value in spec.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_table_borders(table, color=MID_GRAY, *, outside=True, inside_h=True, inside_v=False):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    edge_names = ["top", "left", "bottom", "right", "insideH", "insideV"]
    for edge_name in edge_names:
        edge = tbl_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tbl_borders.append(edge)
        enabled = (
            (edge_name in {"top", "left", "bottom", "right"} and outside)
            or (edge_name == "insideH" and inside_h)
            or (edge_name == "insideV" and inside_v)
        )
        edge.set(qn("w:val"), "single" if enabled else "nil")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border_bottom(paragraph, color=GOLD, size=10, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def add_field(paragraph, field_code: str, display_text: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def style_table_text(table, *, header=True, font_size=8.5, first_col_bold=False, alignments=None):
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0 and header:
                set_cell_shading(cell, NAVY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if alignments and col_idx < len(alignments):
                    paragraph.alignment = alignments[col_idx]
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=font_size,
                        bold=(row_idx == 0 and header) or (first_col_bold and col_idx == 0),
                        color=WHITE if row_idx == 0 and header else INK,
                    )
    if header:
        set_repeat_table_header(table.rows[0])


def add_data_table(doc, headers, rows, weights, *, font_size=8.5, first_col_bold=False, alignments=None, total_rows=None, cell_margins_dxa=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = str(header)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = "" if value is None else str(value)
    widths = column_widths_from_weights(weights, CONTENT_WIDTH_DXA)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=cell_margins_dxa,
    )
    set_table_borders(table, color=MID_GRAY, outside=True, inside_h=True, inside_v=False)
    style_table_text(table, header=True, font_size=font_size, first_col_bold=first_col_bold, alignments=alignments)
    if total_rows:
        for row_idx in total_rows:
            if row_idx < 0:
                row_idx = len(table.rows) + row_idx
            for cell in table.rows[row_idx].cells:
                set_cell_shading(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    return table


def add_callout(doc, label, body, *, fill=LIGHT_BLUE, accent=TEAL, body_color=INK):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, left={"val": "single", "sz": "22", "color": accent, "space": "0"})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.2, bold=True, color=accent)
    r = p.add_run(body)
    set_run_font(r, size=10.2, color=body_color)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=180, cell_margins_dxa={"top": 130, "bottom": 130, "start": 180, "end": 180})
    set_table_borders(table, outside=False, inside_h=False, inside_v=False)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_kpi_strip(doc, metrics):
    table = doc.add_table(rows=2, cols=len(metrics))
    widths = column_widths_from_weights([1] * len(metrics), CONTENT_WIDTH_DXA)
    apply_table_geometry(table, widths, table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa={"top": 110, "bottom": 110, "start": 120, "end": 120})
    set_table_borders(table, color=MID_GRAY, outside=True, inside_h=False, inside_v=True)
    for idx, (value, label) in enumerate(metrics):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        set_cell_shading(table.cell(1, idx), LIGHT_BLUE)
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=17, bold=True, color=NAVY)
        p = table.cell(1, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=8.5, bold=True, color=MUTED)
        table.cell(0, idx).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.cell(1, idx).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_source_note(doc, text):
    p = doc.add_paragraph(style="Source Note")
    p.add_run(text)
    return p


def add_para(doc, text="", *, bold_lead=None, italic=False, align=None, color=None, size=None, keep_with_next=False):
    p = doc.add_paragraph(style="Normal")
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_with_next = keep_with_next
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=color or INK, size=size)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, italic=italic, color=color or INK, size=size)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color or INK, size=size)
    return p


def add_picture(doc, path, *, width=6.35, alt_text="", caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = bool(caption)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    if alt_text:
        shape._inline.docPr.set("descr", alt_text)
        shape._inline.docPr.set("title", alt_text[:100])
    if caption:
        cp = doc.add_paragraph(style="Caption")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.keep_with_next = False
        cp.add_run(caption)
    return shape


def fmt_int(value):
    return f"{value:,.0f}"


def fmt_gbp(value, decimals=0):
    if value < 0:
        return f"(£{abs(value):,.{decimals}f})"
    return f"£{value:,.{decimals}f}"


def fmt_gbp_m(value, decimals=2):
    if value < 0:
        return f"(£{abs(value) / 1_000_000:.{decimals}f}m)"
    return f"£{value / 1_000_000:.{decimals}f}m"


def fmt_gbp_scale(value):
    av = abs(value)
    if av >= 1_000_000_000:
        s = f"£{av / 1_000_000_000:.3g}bn"
    elif av >= 1_000_000:
        s = f"£{av / 1_000_000:.3g}m"
    elif av >= 1_000:
        s = f"£{av / 1_000:.3g}k"
    else:
        s = f"£{av:,.0f}"
    return f"({s})" if value < 0 else s


def fmt_pct(value, decimals=1):
    return f"{value * 100:.{decimals}f}%"


def _chart_fonts():
    font_dir = Path(r"C:\Windows\Fonts")
    regular_path = font_dir / "arial.ttf"
    bold_path = font_dir / "arialbd.ttf"
    return {
        "title": ImageFont.truetype(str(bold_path), 42),
        "axis": ImageFont.truetype(str(regular_path), 27),
        "axis_bold": ImageFont.truetype(str(bold_path), 27),
        "small": ImageFont.truetype(str(regular_path), 23),
        "small_bold": ImageFont.truetype(str(bold_path), 23),
        "label": ImageFont.truetype(str(bold_path), 25),
    }


def _chart_canvas(title):
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    fonts = _chart_fonts()
    draw.text((70, 48), title, font=fonts["title"], fill=f"#{NAVY}")
    draw.line((70, 108, 1730, 108), fill=f"#{GOLD}", width=5)
    return image, draw, fonts


def _line_chart(path, title, labels, series, y_min, y_max, y_step, y_title):
    image, draw, fonts = _chart_canvas(title)
    left, top, right, bottom = 190, 165, 1700, 735
    plot_w, plot_h = right - left, bottom - top
    ticks = []
    value = y_min
    while value <= y_max + 1e-9:
        ticks.append(value)
        value += y_step
    def ypix(v):
        return bottom - (v - y_min) / (y_max - y_min) * plot_h
    for tick in ticks:
        y = ypix(tick)
        draw.line((left, y, right, y), fill="#E2E7ED", width=2)
        draw.text((left - 25, y), f"{tick:g}", font=fonts["small"], fill=f"#{MUTED}", anchor="rm")
    draw.line((left, top, left, bottom), fill="#B8C2CE", width=2)
    draw.line((left, bottom, right, bottom), fill="#B8C2CE", width=2)
    if y_min < 0 < y_max:
        draw.line((left, ypix(0), right, ypix(0)), fill="#8E99A7", width=3)
    xs = [left + i * plot_w / (len(labels) - 1) for i in range(len(labels))]
    for x, label in zip(xs, labels):
        draw.text((x, bottom + 25), label, font=fonts["axis_bold"], fill=f"#{MUTED}", anchor="ma")
    # Keep the unit label horizontal and inside the chart canvas so it remains
    # legible when Word scales the image for the report page.
    draw.text((left, top - 34), y_title, font=fonts["small_bold"], fill="#445268", anchor="la")
    for item in series:
        vals = item["values"]
        points = [(xs[i], ypix(vals[i])) for i in range(len(vals))]
        draw.line(points, fill=item["color"], width=8, joint="curve")
        for idx, (x, y) in enumerate(points):
            draw.ellipse((x - 10, y - 10, x + 10, y + 10), fill=item["color"], outline="white", width=3)
            if item.get("data_labels", True):
                offset = item.get("label_offset", -24)
                draw.text((x, y + offset), item.get("format", lambda v: f"{v:.1f}")(vals[idx]), font=fonts["small_bold"], fill=item["color"], anchor="mm")
    lx, ly = left, 815
    for item in series:
        draw.line((lx, ly, lx + 50, ly), fill=item["color"], width=8)
        draw.text((lx + 66, ly), item["label"], font=fonts["axis"], fill="#445268", anchor="lm")
        lx += draw.textlength(item["label"], font=fonts["axis"]) + 150
    image.save(path)


def save_charts(data):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    years = [y.replace("Year ", "Y") for y in data["operating"]["years"]]
    revenue = [v / 1_000_000 for v in data["operating"]["total_revenue"]]
    ebitda = [v / 1_000_000 for v in data["operating"]["ebitda_pre_one_off"]]
    _line_chart(
        ASSET_DIR / "operating-trajectory.png",
        "Base-case revenue and EBITDA trajectory",
        years,
        [
            {"label": "Revenue", "color": f"#{NAVY}", "values": revenue, "label_offset": -28},
            {"label": "EBITDA before one-offs", "color": f"#{GOLD}", "values": ebitda, "label_offset": 28},
        ],
        -2, 20, 4, "GBP millions",
    )

    image, draw, fonts = _chart_canvas("Digital document mix expands while paper becomes a fallback")
    left, top, right, bottom = 180, 165, 1700, 700
    plot_w, plot_h = right - left, bottom - top
    for tick in [0, 25, 50, 75, 100]:
        y = bottom - tick / 100 * plot_h
        draw.line((left, y, right, y), fill="#E2E7ED", width=2)
        draw.text((left - 25, y), f"{tick}%", font=fonts["small"], fill=f"#{MUTED}", anchor="rm")
    draw.line((left, top, left, bottom), fill="#B8C2CE", width=2)
    draw.line((left, bottom, right, bottom), fill="#B8C2CE", width=2)
    mix = data["tier_mix"]
    centers = [left + (i + 0.5) * plot_w / 5 for i in range(5)]
    bar_w = 150
    colors = [f"#{TEAL}", "#6FA9A4", f"#{GOLD}"]
    mix_series = [mix["tier_a"], mix["tier_b"], mix["tier_c"]]
    for idx, x in enumerate(centers):
        cumulative = 0
        for values, color in zip(mix_series, colors):
            value = values[idx] * 100
            y1 = bottom - cumulative / 100 * plot_h
            cumulative += value
            y2 = bottom - cumulative / 100 * plot_h
            draw.rectangle((x - bar_w / 2, y2, x + bar_w / 2, y1), fill=color)
            if value >= 8:
                draw.text((x, (y1 + y2) / 2), f"{value:.0f}%", font=fonts["small_bold"], fill="white", anchor="mm")
        draw.text((x, bottom + 24), years[idx], font=fonts["axis_bold"], fill=f"#{MUTED}", anchor="ma")
    gm = [v * 100 for v in data["operating"]["gross_margin"]]
    gm_points = [(centers[i], bottom - gm[i] / 100 * plot_h) for i in range(5)]
    draw.line(gm_points, fill=f"#{NAVY}", width=8, joint="curve")
    for x, y in gm_points:
        draw.rectangle((x - 9, y - 9, x + 9, y + 9), fill=f"#{NAVY}", outline="white", width=3)
    legends = [(colors[0], "Tier A - eBL"), (colors[1], "Tier B - carrier/API"), (colors[2], "Tier C - paper"), (f"#{NAVY}", "Company gross margin")]
    lx, ly = left, 810
    for color, label in legends:
        draw.rectangle((lx, ly - 10, lx + 36, ly + 10), fill=color)
        draw.text((lx + 50, ly), label, font=fonts["axis"], fill="#445268", anchor="lm")
        lx += draw.textlength(label, font=fonts["axis"]) + 115
    image.save(ASSET_DIR / "tier-mix-and-margin.png")

    image, draw, fonts = _chart_canvas("Initial operating funding use")
    use = data["use_of_funds"]
    left, top, right, bottom = 470, 165, 1660, 760
    max_amount = max(u["amount"] for u in use) / 1_000_000 * 1.28
    bar_h = 78
    gap = 27
    palette = [f"#{NAVY}", "#54718A", f"#{TEAL}", f"#{GOLD}", "#8C9AA8"]
    for idx, (item, color) in enumerate(zip(use, palette)):
        y = top + idx * (bar_h + gap)
        label = item["category"].replace("Compliance, legal & regulatory", "Compliance, legal & regulatory")
        draw.text((left - 30, y + bar_h / 2), label, font=fonts["axis_bold"], fill="#445268", anchor="rm")
        width = (item["amount"] / 1_000_000) / max_amount * (right - left)
        draw.rectangle((left, y, left + width, y + bar_h), fill=color)
        draw.text((left + width + 24, y + bar_h / 2), f"£{item['amount']/1_000_000:.2f}m | {item['share']*100:.1f}%", font=fonts["axis_bold"], fill="#445268", anchor="lm")
    image.save(ASSET_DIR / "use-of-funds.png")

    checkpoints = data["cash_runway"]["checkpoints"]
    cash_series = [
        {"label": "Low", "color": "#9FA8B3", "values": [v / 1_000_000 for v in data["cash_runway"]["low"]], "data_labels": False},
        {"label": "Base", "color": f"#{TEAL}", "values": [v / 1_000_000 for v in data["cash_runway"]["base"]], "data_labels": False},
        {"label": "High", "color": f"#{GOLD}", "values": [v / 1_000_000 for v in data["cash_runway"]["high"]], "data_labels": False},
    ]
    _line_chart(
        ASSET_DIR / "cash-runway-scenarios.png",
        "Cash runway diverges sharply after Year 3",
        checkpoints,
        cash_series,
        -5, 35, 5, "Cumulative operating cash flow, GBP m",
    )


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = rgb(NAVY)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.widow_control = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(TEAL)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(DEEP_NAVY)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)

    source = doc.styles.add_style("Source Note", 1)
    source.font.name = "Calibri"
    source._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    source._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    source.font.size = Pt(8)
    source.font.italic = True
    source.font.color.rgb = rgb(MUTED)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = "Blockmediary: Financial Value and Commercial Viability Report"
    doc.core_properties.subject = "Unit economics, five-year operating case, capital requirement and commercial validation"
    doc.core_properties.author = "Team Transakt"
    doc.core_properties.keywords = "Blockmediary, documentary escrow, financial model, unit economics, commercial viability"

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run("BLOCKMEDIARY  |  FINANCIAL VALUE")
    set_run_font(r, size=8.2, bold=True, color=NAVY, all_caps=True)
    r = hp.add_run("\tTEAM TRANSAKT")
    set_run_font(r, size=8.2, bold=True, color=MUTED, all_caps=True)
    set_paragraph_border_bottom(hp, color=GOLD, size=6, space=4)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run("Team Transakt | Blockmediary | 13 August 2026")
    set_run_font(r, size=8, color=MUTED)
    fp.add_run("\t")
    r = fp.add_run("Page ")
    set_run_font(r, size=8, color=MUTED)
    r = add_field(fp, "PAGE", "1")
    set_run_font(r, size=8, color=MUTED)
    r = fp.add_run(" of ")
    set_run_font(r, size=8, color=MUTED)
    r = add_field(fp, "NUMPAGES", "12")
    set_run_font(r, size=8, color=MUTED)

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ffp.add_run("TEAM TRANSAKT  |  BLOCKMEDIARY  |  FINANCIAL REPORT")
    set_run_font(r, size=8.2, bold=True, color=MUTED, all_caps=True)


def page_break(doc):
    doc.add_page_break()


def add_section_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    set_paragraph_border_bottom(p, color=GOLD, size=5, space=4)
    return p


def build_report(data):
    doc = Document()
    configure_document(doc)
    op = data["operating"]
    timing = data["timing_and_capital"]
    base_scenario = next(s for s in data["scenarios"] if s["name"] == "Base")
    low_scenario = next(s for s in data["scenarios"] if s["name"] == "Low")

    # Cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(54)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    r = kicker.add_run("BLOCKMEDIARY  |  FINANCIAL REPORT")
    set_run_font(r, size=9, bold=True, color=GOLD, all_caps=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(9)
    r = title.add_run("Financial Value and\nCommercial Viability")
    set_run_font(r, size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(15)
    r = subtitle.add_run("Unit economics, five-year operating case, capital requirement and commercial validation")
    set_run_font(r, size=13.5, color=TEAL)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(20)
    set_paragraph_border_bottom(rule, color=GOLD, size=12, space=1)

    add_callout(
        doc,
        "Investment view",
        "Each workflow tier produces positive transaction contribution before shared platform costs. Commercial viability still depends on distribution, regulatory execution and production controls.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )
    add_kpi_strip(doc, [
        (fmt_gbp_m(timing["initial_operating_funding_need"], 2), "Initial operating funding"),
        (f"Month {timing['first_paid_pilot_month']}", "First paid pilot"),
        (fmt_gbp_m(op["total_revenue"][2], 2), "Base Year-3 revenue"),
        (base_scenario["monthly_break_even"], "Monthly cash-flow break-even"),
    ])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Prepared by Team Transakt")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Preliminary commercial planning report")
    set_run_font(r, size=9.5, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Reference model: Blockmediary Financial Model, 13 August 2026")
    set_run_font(r, size=9.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Status: preliminary planning report; not an audited forecast, valuation or investment recommendation")
    set_run_font(r, size=9.5, color=MUTED)

    page_break(doc)

    # Executive conclusion
    add_section_heading(doc, "Executive conclusion")
    add_para(doc, "Blockmediary charges for prefunded, document-conditioned settlement. At the model's Year-1 deal sizes, the standard fee is £305 for a £35,000 Tier A eBL workflow, £540 for a £30,000 Tier B carrier/API workflow and £1,200 for a £30,000 Tier C paper workflow. Those fees equal 0.9%, 1.8% and 4.0% of deal value respectively, before any conditional dispute fee. Customers pay for payment assurance and document verification. Blockmediary does not lend, take FX risk or earn a spread on client money.")
    add_para(doc, "The base case reaches £76.0m of Year-3 gross merchandise value (GMV) across 2,000 deals. It converts that throughput into £1.371m of revenue, £1.057m of gross profit and a 77.1% gross margin, but still records negative EBITDA of £409k. Monthly cash-flow break-even occurs in Year 4, Month 8, after annual deal volume rises from 2,000 to 15,000 and active customers rise from 333 to 1,875. Transaction margins are positive before company profitability because payroll, fixed operating cost and regulatory overhead require much greater distribution scale.")
    add_callout(
        doc,
        "Financial judgement",
        f"Capital should be committed in stages. The initial operating requirement is {fmt_gbp(timing['initial_operating_funding_need'])}; the low case needs {fmt_gbp(low_scenario['funding_need'])} and remains below monthly break-even beyond Month 60.",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    doc.add_paragraph("Basis of analysis", style="Heading 2")
    basis_rows = [
        ["Revenue model", "Escrow, document, exception and partner/API fees", "No lending margin, FX spread or client-fund yield"],
        ["Planning horizon", "60 monthly periods with five annual summaries", "Pre-tax and pre-financing operating case"],
        ["Scenarios", "Low, Base and High deal volumes and tier mix", "Payroll, fixed opex and one-offs remain fixed"],
        ["Capital treatment", "Peak operating gap plus a liquidity reserve", "VARA restricted capital is separate from runway"],
        ["Current stage", "Base Sepolia proof of concept using test USDC", "No live revenue, customers or permission to handle real funds"],
    ]
    add_data_table(
        doc,
        ["Area", "Model treatment", "Boundary"],
        basis_rows,
        [1.55, 2.45, 2.5],
        font_size=8.2,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] * 3,
    )
    add_source_note(doc, "Source: Blockmediary Financial Model, README, Assumptions, Launch Readiness, Scenario Engine and Client Funds Memo.")

    page_break(doc)

    # Real-world value
    add_section_heading(doc, "1. Commercial value by stakeholder")
    add_para(doc, "The buyer prefunds the agreed stablecoin amount, and the seller can verify that funds are locked before shipping. Release follows documented compliance and the objection process. Blockmediary charges for administering that workflow without advancing capital against the trade.")
    value_rows = [
        ["Buyer", "Avoid paying the seller outright before shipment evidence exists", "Funds remain locked pending document conditions", "The service fee is small relative to the trade principal, but it is not an insurance guarantee"],
        ["Seller", "Avoid shipping with only an unsecured promise of later payment", "Prefunding is visible before shipment; valid release does not depend on discretionary buyer approval", "Economic value is faster, more certain access to agreed funds when documents comply"],
        ["Blockmediary", "Monetise trust without lending or balance-sheet credit exposure", "Escrow, document, exception and partner/API fees", "Base total revenue yield falls from 2.7% to 1.2% of GMV as scale and digital mix increase"],
        ["Partner / regulator", "Operate a controlled and auditable settlement perimeter", "Client assets are gross memorandum balances with matching liabilities", "0% is available to operations and no client-fund yield is recognised"],
    ]
    add_data_table(
        doc,
        ["Stakeholder", "Economic problem", "Financial mechanism", "Value interpretation"],
        value_rows,
        [1.0, 2.0, 1.85, 1.65],
        font_size=8.0,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] * 4,
    )
    add_source_note(doc, "Sources: Financial Model, P&L 5yr!C7:G31 and Client Funds Memo!A5:H31; Trade Escrow Agreement, pp.2-4 and 18-19.")

    doc.add_paragraph("How much trade value is supported by each pound of revenue?", style="Heading 2")
    add_para(doc, f"In the base case, the platform processes {op['gmv_per_revenue_pound'][2]:.1f} pounds of Year-3 GMV for every pound of company revenue. The multiple expands to {op['gmv_per_revenue_pound'][4]:.1f}x by Year 5 because the digital tier gains share and the blended escrow take rate declines. GMV belongs to customers and is never company revenue or operating cash.")
    add_kpi_strip(doc, [
        (f"{op['gmv_per_revenue_pound'][2]:.1f}x", "Y3 GMV / revenue"),
        (fmt_pct(op["revenue_yield"][2]), "Y3 total revenue yield"),
        (fmt_pct(op["gross_margin"][2]), "Y3 company gross margin"),
        (fmt_gbp_scale(data["client_funds"]["average_safeguarded_assets"][2]), "Y3 avg safeguarded assets"),
    ])
    add_callout(doc, "Boundary", "GMV, client escrow assets and future VARA restricted capital are not company cash. The model gives all three zero operating availability.", fill=LIGHT_RED, accent=RED)

    page_break(doc)

    # Transaction economics
    add_section_heading(doc, "2. Economics of one transaction")
    add_callout(doc, "Price formula", "Standard fee = max(deal value x tier take rate, minimum escrow fee) + document/review fee. Conditional dispute or amendment fees are excluded until the event occurs.", fill=LIGHT_BLUE, accent=NAVY)
    pricing_rows = []
    for tier in data["tiers"]:
        workflow = tier["name"].split(" - ", 1)[1]
        pricing_rows.append([
            tier["name"].split(" - ", 1)[0],
            workflow,
            fmt_gbp(tier["deal_value"]),
            fmt_pct(tier["take_rate"]),
            fmt_gbp(tier["standard_fee"]),
            fmt_pct(tier["standard_fee_pct_of_deal"]),
            fmt_gbp(tier["conditional_dispute_fee"]),
        ])
    add_data_table(
        doc,
        ["Tier", "Workflow", "Y1 deal", "Take rate", "Standard fee", "Fee / deal", "Conditional event fee"],
        pricing_rows,
        [0.65, 1.35, 1.0, 0.75, 1.0, 0.85, 1.2],
        font_size=8.1,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 5,
    )
    add_source_note(doc, "Source: Financial Model, Assumptions!F63:F74 and F149:F154; Tier Economics!B3:U18. Standard fee excludes any dispute/amendment event.")

    doc.add_paragraph("Contribution economics before shared platform costs", style="Heading 2")
    contribution_rows = []
    for tier in data["tiers"]:
        contribution_rows.append([
            tier["name"].split(" - ", 1)[0],
            fmt_gbp(tier["expected_revenue_per_deal"]),
            fmt_gbp(tier["variable_cost_per_deal"], 1),
            fmt_gbp(tier["contribution_per_deal_before_shared_cogs"], 1),
            fmt_pct(tier["tier_gross_margin_before_shared_cogs"]),
        ])
    add_data_table(
        doc,
        ["Tier", "Expected revenue / deal", "Tier-specific variable cost", "Contribution / deal", "Margin before shared COGS"],
        contribution_rows,
        [0.8, 1.55, 1.7, 1.55, 1.5],
        font_size=8.3,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 4,
    )
    add_source_note(doc, "Source: Financial Model, Tier Economics!B15:D26. Expected revenue includes probability-weighted dispute/amendment fees; costs exclude shared COGS, payroll and fixed overhead.")

    add_para(doc, "Tier A has the lowest cost to serve. At the Year-1 assumptions, a £305 standard customer fee supports a £35,000 trade. Expected revenue of £309 and tier-specific variable cost of £10.40 produce a 96.6% margin before shared platform cost. Year-1 company gross margin is lower at 71.1% after shared COGS, and EBITDA remains negative after payroll and fixed operating expenditure.", bold_lead="Tier A has the lowest cost to serve.")
    add_para(doc, "Tier B requires more document and human/API handling. It retains a 91.4% margin before shared COGS. Its £540 standard fee equals 1.8% of the modeled £30,000 deal value.", bold_lead="Tier B requires more document and human/API handling.")
    add_para(doc, "Tier C prices for manual paper processing. It has the highest fee, expected dispute rate and cost to serve. Its 70.0% margin before shared COGS remains positive but is materially below the digital tiers. The scale plan reduces the paper mix over time.", bold_lead="Tier C prices for manual paper processing.")
    add_callout(doc, "Validation requirement", "These prices are management hypotheses informed by the deal-value and competitor reports. They are not an observed tariff or proof of willingness to pay. The paid pilot must test conversion by tier, deal value and fee.", fill=LIGHT_GOLD, accent=GOLD)

    page_break(doc)

    # Operating leverage
    add_section_heading(doc, "3. Digital mix creates operating leverage")
    add_picture(
        doc,
        ASSET_DIR / "tier-mix-and-margin.png",
        width=6.35,
        alt_text="Stacked bars show Tier A increasing from 15% to 70%, Tier C falling from 40% to 5%, and company gross margin remaining around 71% to 79% across five years.",
        caption="Figure 1. Base-case tier mix and company gross margin.",
    )
    add_source_note(doc, "Source: Financial Model, Tier Economics!B3:U26 and P&L 5yr!C29:G31.")
    mix = data["tier_mix"]
    leverage_rows = [
        ["Tier A share", fmt_pct(mix["tier_a"][0], 0), fmt_pct(mix["tier_a"][2], 0), fmt_pct(mix["tier_a"][4], 0), "eBL becomes the dominant scale path"],
        ["Tier C share", fmt_pct(mix["tier_c"][0], 0), fmt_pct(mix["tier_c"][2], 0), fmt_pct(mix["tier_c"][4], 0), "paper remains available but becomes exceptional"],
        ["Blended escrow take rate", fmt_pct(mix["blended_escrow_take_rate"][0]), fmt_pct(mix["blended_escrow_take_rate"][2]), fmt_pct(mix["blended_escrow_take_rate"][4]), "lower pricing accompanies digitisation"],
        ["Documents per deal", f"{mix['docs_per_deal'][0]:.2f}", f"{mix['docs_per_deal'][2]:.2f}", f"{mix['docs_per_deal'][4]:.2f}", "weighted handling load declines"],
        ["Margin before shared COGS", fmt_pct(mix["tier_margin_before_shared_cogs"][0]), fmt_pct(mix["tier_margin_before_shared_cogs"][2]), fmt_pct(mix["tier_margin_before_shared_cogs"][4]), "mix and automation improve contribution economics"],
        ["Company gross margin", fmt_pct(op["gross_margin"][0]), fmt_pct(op["gross_margin"][2]), fmt_pct(op["gross_margin"][4]), "shared risk and operating costs keep the margin realistic"],
    ]
    add_data_table(
        doc,
        ["Metric", "Year 1", "Year 3", "Year 5", "Commercial meaning"],
        leverage_rows,
        [1.55, 0.75, 0.75, 0.75, 2.7],
        font_size=8.2,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_source_note(doc, "Source: Financial Model, Tier Economics!E4:U26 and P&L 5yr!C31:G31.")
    add_para(doc, "The digital document pathway costs less to operate than paper review. As digital tiers gain share, the blended escrow take rate falls from approximately 2.0% to 1.0% while whole-company gross margin stays within 71% to 79%. The forecast therefore depends on the modeled eBL/API adoption curve.")

    page_break(doc)

    # Five-year operating case
    add_section_heading(doc, "4. Five-year operating case")
    add_picture(
        doc,
        ASSET_DIR / "operating-trajectory.png",
        width=6.35,
        alt_text="Line chart shows base-case revenue rising from near zero in Year 1 to 17.8 million pounds in Year 5, while EBITDA remains negative through Year 3 and turns positive in Year 4.",
        caption="Figure 2. Base-case revenue and EBITDA trajectory.",
    )
    add_source_note(doc, "Source: Financial Model, P&L 5yr!C18:G18 and C61:G61.")

    scale_rows = []
    for idx, year in enumerate(op["years"]):
        scale_rows.append([
            year,
            fmt_int(op["deals"][idx]),
            fmt_int(op["active_customers"][idx]),
            fmt_gbp(op["avg_deal_value"][idx]),
            fmt_gbp_scale(op["gmv"][idx]),
            fmt_gbp_scale(op["total_revenue"][idx]),
        ])
    add_data_table(
        doc,
        ["Year", "Deals", "Active customers", "Avg deal", "GMV", "Revenue"],
        scale_rows,
        [0.75, 0.75, 1.2, 1.1, 1.25, 1.25],
        font_size=8.4,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 5,
    )
    add_source_note(doc, "Source: Financial Model, P&L 5yr!C7:G18 and C69:G84. Active customers are model-implied targets, rounded for presentation.")

    profitability_rows = []
    for idx, year in enumerate(op["years"]):
        profitability_rows.append([
            year,
            fmt_gbp_scale(op["gross_profit"][idx]),
            fmt_pct(op["gross_margin"][idx]),
            fmt_gbp_scale(op["ebitda_pre_one_off"][idx]),
            fmt_gbp_scale(op["operating_result"][idx]),
            fmt_gbp_scale(op["cumulative_result"][idx]),
        ])
    add_data_table(
        doc,
        ["Year", "Gross profit", "Gross margin", "EBITDA before one-offs", "Operating result", "Cumulative result"],
        profitability_rows,
        [0.7, 1.05, 0.9, 1.4, 1.2, 1.25],
        font_size=8.2,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 5,
    )
    add_source_note(doc, "Source: Financial Model, P&L 5yr!C29:G66. The model is pre-tax and pre-financing; corporation tax, interest and depreciation/amortisation are not modeled.")

    page_break(doc)

    doc.add_paragraph("The Year-3 bridge: proof of monetisation, not yet profitability", style="Heading 2")
    rb = data["revenue_breakdown"]
    y3 = 2
    bridge_items = [
        ("Escrow transaction fees", rb["escrow_transaction_fees"][y3]),
        ("Launch discounts and credits", rb["launch_discounts"][y3]),
        ("Document and exception-handling fees", rb["document_and_exception_fees"][y3]),
        ("Partner / API and ancillary services", rb["partner_api_and_ancillary"][y3]),
        ("Total revenue", rb["total_revenue"][y3]),
    ]
    bridge_rows = [[name, fmt_gbp(amount), fmt_pct(amount / rb["total_revenue"][y3]) if name != "Total revenue" else "100.0%"] for name, amount in bridge_items]
    add_data_table(
        doc,
        ["Year-3 revenue stream", "Amount", "Share of total"],
        bridge_rows,
        [3.7, 1.4, 1.4],
        font_size=8.7,
        first_col_bold=False,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_rows=[-1],
    )
    add_source_note(doc, "Source: Financial Model, P&L 5yr!E14:E18.")
    add_para(doc, "Escrow fees provide 74.6% of Year-3 revenue. Document and exception fees provide 17.1%, and partner/API services provide 13.5%, partly offset by launch discounts and credits. Gross profit of £1.057m does not cover £1.015m of payroll and £451k of fixed operating expenditure. The business needs more deals, more revenue per active customer or slower fixed-cost growth before it becomes self-funding.")

    doc.add_paragraph("The Year-4 inflection is the largest assumption in the model", style="Heading 2")
    deal_growth = op["deals"][3] / op["deals"][2]
    customer_growth = op["active_customers"][3] / op["active_customers"][2]
    revenue_growth = op["total_revenue"][3] / op["total_revenue"][2]
    add_callout(
        doc,
        "Scale test",
        f"From Year 3 to Year 4, deals rise {deal_growth:.1f}x, active customers rise {customer_growth:.1f}x and revenue rises {revenue_growth:.1f}x. Gross new customers required per Sales/BD/CS FTE rise from {op['new_customers_per_commercial_fte'][2]:.0f} to {op['new_customers_per_commercial_fte'][3]:.0f}, against a model warning threshold of 250.",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )
    add_para(doc, "The commercial capacity check remains within its modeled threshold, but it does not validate demand. The Year-4 and Year-5 outputs depend on partner distribution, self-serve onboarding, digital-document adoption and an AI-enabled operating model. They are conditional operating targets, not a straight-line forecast from the prototype.")

    page_break(doc)

    # Capital
    add_section_heading(doc, "5. Capital requirement and use of funds")
    add_para(doc, f"The base case requires {fmt_gbp(timing['initial_operating_funding_need'])} of initial operating capital. This equals the modeled peak monthly operating deficit of approximately {fmt_gbp(base_scenario['peak_cash_gap'])} plus a three-month operating reserve of {fmt_gbp(op['reserve'][-1])}. Future VARA restricted capital of {fmt_gbp(timing['vara_restricted_capital'])} is separate and cannot be used as runway.")
    add_picture(
        doc,
        ASSET_DIR / "use-of-funds.png",
        width=6.3,
        alt_text="Horizontal bars allocate the 3.063 million pound funding need: 34.8% compliance and regulatory, 34.7% operations and people, 18.6% engineering, 6.3% reserve, and 5.6% go-to-market.",
        caption="Figure 3. Initial operating funding allocation.",
    )
    add_source_note(doc, "Source: Financial Model, Dashboard!J57:M62.")

    capital_rows = [
        ["Peak operating cash gap", fmt_gbp(base_scenario["peak_cash_gap"]), "Maximum cumulative deficit before financing"],
        ["Three-month operating reserve", fmt_gbp(op["reserve"][-1]), "Included in initial operating funding"],
        ["Initial operating funding", fmt_gbp(timing["initial_operating_funding_need"]), "Spendable operating runway"],
        ["Future VARA restricted capital", fmt_gbp(timing["vara_restricted_capital"]), "Separate; not operating cash"],
        ["Combined staged capital envelope", fmt_gbp(timing["combined_staged_capital_envelope"]), "Derived total; not the initial ask"],
    ]
    add_data_table(
        doc,
        ["Capital item", "Amount", "Treatment"],
        capital_rows,
        [2.65, 1.25, 2.6],
        font_size=8.6,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
        total_rows=[2],
    )
    add_source_note(doc, "Source: Financial Model, Launch Readiness!E12:E17; P&L 5yr!F64:F66; Dashboard!J57:M62.")

    # Keep the regulatory-cost evidence together instead of splitting the
    # table across two pages beneath the use-of-funds analysis.
    page_break(doc)
    doc.add_paragraph("What the capital is buying", style="Heading 2")
    add_para(doc, "Operations, people and G&A account for 34.7% of the operating requirement. Compliance, legal and regulatory work account for a further 34.8%. Engineering and product are 18.6%, go-to-market is 5.6%, and reserve/contingency is 6.3%. The prototype reduces product-feasibility uncertainty, but production security, key management, partner integration and an independent audit remain outstanding.")
    regulatory_rows = [
        ["First-six-month startup build", fmt_gbp(timing["startup_six_month_total_including_contingency"]), "Includes 15% contingency"],
        ["DIFC / regulated-partner pilot onboarding", fmt_gbp(timing["partner_pilot_one_off"]), "Year-1 one-off"],
        ["VARA application placeholder", fmt_gbp(timing["vara_application_fee_gbp"]), "Year-1 staged application cost"],
        ["Total Year-1 startup and staged regulatory one-offs", fmt_gbp(timing["startup_six_month_total_including_contingency"] + timing["partner_pilot_one_off"] + timing["vara_application_fee_gbp"]), "Modeled operating expense"],
        ["Partner-pilot oversight", fmt_gbp(timing["partner_pilot_annual_run_rate"]), "Annual pilot-stage run-rate"],
        ["VARA supervision placeholder", fmt_gbp(timing["vara_supervision_gbp_per_year"]), "Annual cost after licensing"],
    ]
    add_data_table(
        doc,
        ["Cost item", "Model value", "Financial treatment"],
        regulatory_rows,
        [3.1, 1.25, 2.15],
        font_size=8.4,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
        total_rows=[4],
    )
    add_source_note(doc, "Sources: Financial Model, Startup 6mo!F40:F44 and Launch Readiness!A5:E17. VARA figures are planning placeholders pending counsel and regulator confirmation.")

    doc.add_paragraph("Reserve sensitivity", style="Heading 2")
    reserve_rows = [[f"{item['months']} months", item["reserve"], item["funding_need"]] for item in data["reserve_sensitivity"]]
    add_data_table(
        doc,
        ["Reserve policy", "Operating reserve", "Funding need"],
        reserve_rows,
        [2.3, 1.9, 2.3],
        font_size=8.7,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
    )
    add_source_note(doc, "Source: Financial Model, P&L 5yr!C89:E91. Rounded presentation values.")

    page_break(doc)

    # Scenarios
    add_section_heading(doc, "6. Downside, base and upside")
    add_picture(
        doc,
        ASSET_DIR / "cash-runway-scenarios.png",
        width=6.35,
        alt_text="Line chart shows cumulative operating cash flow through Year 5: the low case stays negative, the base case turns positive during Year 4, and the high case rises strongly after Year 3.",
        caption="Figure 4. Six-month cumulative operating cash-flow checkpoints by scenario.",
    )
    add_source_note(doc, "Source: Financial Model, Dashboard!B67:E77 and Cash Flow Statement!B45:F47.")

    scenario_operating = []
    scenario_cash = []
    for scenario in data["scenarios"]:
        scenario_operating.append([
            scenario["name"],
            fmt_gbp_scale(scenario["year_3_revenue"]),
            fmt_gbp_scale(scenario["year_3_ebitda"]),
            fmt_int(scenario["year_3_active_customers"]),
            scenario["monthly_break_even"],
        ])
        scenario_cash.append([
            scenario["name"],
            fmt_gbp_m(scenario["peak_cash_gap"], 3),
            fmt_gbp_m(scenario["funding_need"], 3),
            fmt_gbp_m(scenario["month_60_operating_cash_flow"], 3),
        ])
    add_data_table(
        doc,
        ["Case", "Y3 revenue", "Y3 EBITDA", "Y3 active customers", "Monthly break-even"],
        scenario_operating,
        [0.85, 1.2, 1.2, 1.45, 1.8],
        font_size=8.5,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_source_note(doc, "Source: Financial Model, Scenario Engine!D30:P62 and Cash Flow Statement!B45:B47.")
    add_data_table(
        doc,
        ["Case", "Peak cash gap", "Initial operating funding", "Month-60 operating cash flow"],
        scenario_cash,
        [1.0, 1.55, 1.85, 2.1],
        font_size=8.5,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 3,
    )
    add_source_note(doc, "Source: Financial Model, Cash Flow Statement!B45:F47 and Dashboard!J23:N28.")

    add_para(doc, "The low case remains cash negative through Month 60. It requires £4.279m of operating funding and ends the five-year period with cumulative operating cash flow of negative £3.104m. Payroll and fixed operating costs do not fall when commercial volume misses.", bold_lead="The low case remains cash negative through Month 60.")
    add_para(doc, "In the base case, monthly break-even occurs in Year 4, Month 8, and cumulative operating cash flow reaches positive £12.717m by Month 60 before financing. The result depends on the Year-4 volume step rather than a higher price assumption.", bold_lead="In the base case, monthly break-even occurs in Year 4, Month 8")
    add_para(doc, "The high case reaches positive Year-3 EBITDA and monthly break-even in Year 4, Month 4, with a lower £2.455m initial operating requirement. It is an upside case rather than the fundraising base.", bold_lead="The high case reaches positive Year-3 EBITDA")
    page_break(doc)

    # Current operating stage
    add_section_heading(doc, "7. Current operating stage and dependencies")
    add_para(doc, "Blockmediary currently runs as a proof of concept on Base Sepolia testnet. It includes a non-custodial escrow contract, structured deal intake, deterministic document-rule grading, a release and refund path, an objection window, multi-deal dashboards and an audit trail. Mainnet deployment, independent security audit, hardware-secured release keys, live KYC/KYB/sanctions screening, AI/OCR extraction and regulatory permissions are not yet in place.")
    stage_rows = [
        ["Settlement", "Escrow create, fund, release and refund on Base Sepolia with test USDC", "Mainnet/token approval, custody/control opinion, production key management"],
        ["Document verification", "Deterministic rules engine; structured bill-of-lading data", "AI/OCR extraction, wider document set, human review console"],
        ["Compliance", "Requirements, risk register and legal template designed", "Live KYB/KYC/sanctions controls, partner sign-off and licensing"],
        ["Commercial", "Pricing, deal values, tier economics and five-year model", "Observed willingness to pay, paid pilot conversion, churn and deal frequency"],
        ["Financial", "42 control checks pass and scenarios reconcile", "Actual revenue, cost-to-serve, holding period and acquisition performance"],
    ]
    add_data_table(
        doc,
        ["Dimension", "Current position", "Required before commercial scale"],
        stage_rows,
        [1.05, 2.65, 2.8],
        font_size=8.2,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] * 3,
    )
    add_source_note(doc, "Sources: repository README and Business Requirements Document; DIFC / VARA Readiness Report, pp.2-3 and 7-11; Legal and Compliance Risk Assessment, p.2.")

    doc.add_paragraph("Commercial and regulatory timeline", style="Heading 2")
    timeline_rows = [
        ["Now", "Working testnet prototype", "Technical feasibility and deterministic release workflow demonstrated"],
        ["Month 1", "DIFC partner-pilot and VARA workstreams begin", "Counsel, partner perimeter and application work start in parallel"],
        [f"Month {timing['first_paid_pilot_month']}", "First paying partner-led pilot", "First external evidence for price, conversion and operating cost"],
        [f"Month {timing['first_cash_receipt_month']}", "First cash receipt", "One-month collection lag modeled"],
        [f"Month {timing['full_vara_licence_month']}", "Full VARA target - base timing", "Wider Dubai scale route, subject to permission scope and approval"],
        [base_scenario["monthly_break_even"], "Monthly cash-flow break-even", "Base case only; depends on the Year-4 distribution step"],
    ]
    add_data_table(
        doc,
        ["Timing", "Milestone", "Financial meaning"],
        timeline_rows,
        [1.0, 2.35, 3.15],
        font_size=8.4,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] * 3,
    )
    add_source_note(doc, "Sources: Financial Model, Launch Readiness!A19:J23 and Cash Flow Statement!B45:B47; DIFC / VARA Readiness Report, pp.1-3 and 11.")
    add_callout(doc, "Current scope", "The testnet escrow and deterministic document-release workflow are operational. Blockmediary has no live revenue, customers or permission to transact with real funds.", fill=LIGHT_TEAL, accent=TEAL)

    page_break(doc)

    # Investor proof requirements
    add_section_heading(doc, "8. What must be proven before scale capital")
    add_para(doc, "Before scale capital is committed, the paid pilot must measure each unobserved driver and replace the corresponding model assumption. The five-year outputs are operating scenarios, not a valuation.")
    proof_rows = [
        ["Deal value distribution", "15-20 target SME interviews plus anonymised invoices; ideally at least 200 recent UAE forwarder jobs", "Median, quartiles, right tail and share above £50k by tier"],
        ["Willingness to pay", "20-30 target-user price tests and paid-pilot conversion", "Conversion probability by tier, deal value and fee"],
        ["Document operations", "Time-and-motion study with documentary reviewers", "Minutes per pack, rework, discrepancy, escalation and loaded cost"],
        ["Commercial motion", "Pilot cohort telemetry", "Acquisition source, CAC, churn, deals per active customer and revenue per customer"],
        ["Settlement operations", "Pilot transaction telemetry", "Funded-to-release days, failure rate, objection rate and support burden"],
        ["Regulatory perimeter", "Integrated DFSA / VARA / CBUAE memorandum plus partner responsibility matrix", "Permission scope, entity design, capital, fees and launch legality"],
    ]
    add_data_table(
        doc,
        ["Driver to validate", "Minimum evidence", "Model input to replace"],
        proof_rows,
        [1.45, 2.7, 2.35],
        font_size=8.1,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] * 3,
    )
    add_source_note(doc, "Sources: Deal Value Research Report, pp.24-25; DIFC / VARA Readiness Report, p.11; Financial Model, Assumptions, Tier Economics and Client Funds Memo.")

    doc.add_paragraph("Recommended financing discipline", style="Heading 2")
    financing_rows = [
        ["Gate 1 - perimeter and production readiness", "Integrated UAE perimeter advice, named regulated partner, responsibility matrix, production security and audit plan", "Confirms that the proposed pilot can legally and safely generate evidence"],
        ["Gate 2 - paid pilot", "Signed pilot customer, measured price acceptance, document workload, funded-to-release time and exception rate", "Replaces the highest-impact unobserved unit-economics assumptions"],
        ["Gate 3 - repeatability", "Repeat customers, partner channel evidence, stable gross margin and conversion by tier", "Tests the Year-3 customer and deal-frequency bridge"],
        ["Gate 4 - scale and licensing", "Application progress, distribution capacity, product automation and control evidence", "Supports the Year-4 volume step and wider licence investment"],
    ]
    add_data_table(
        doc,
        ["Funding gate", "Evidence required", "Why it protects capital"],
        financing_rows,
        [1.45, 3.05, 2.0],
        font_size=8.2,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] * 3,
    )
    add_source_note(doc, "Interpretive recommendation based on the financial model and companion research; the workbook does not prescribe funding tranches.")
    add_callout(doc, "Decision rule", "Do not release scale capital merely because technical milestones are complete. The paid pilot must replace price, workload, conversion and cycle-time assumptions with observed data.", fill=LIGHT_GOLD, accent=GOLD)

    page_break(doc)

    # Model integrity
    add_section_heading(doc, "9. Model integrity, definitions and limitations")
    checks = data["model_checks"]
    add_callout(doc, "Mechanical status", f"PASS: {checks['passed']} of {checks['count']} model checks pass, and the formula-error scan found {len(checks['formula_errors'])} cached #REF!, #DIV/0!, #VALUE!, #NAME?, #N/A or #NUM! errors.", fill=LIGHT_GREEN, accent=GREEN)
    add_para(doc, "A PASS means the workbook reconciles mechanically: the scenario selector, P&L roll-ups, cash roll-forward, revenue and COGS ties, client-fund memorandum, use-of-funds bridge, launch timing and restricted-capital treatment all agree. It does not mean that customers will buy, regulators will approve the route on schedule or Year-4 distribution will occur.")

    definitions = [
        ["GMV", "Customer transaction value processed. It is not revenue, profit or company cash."],
        ["Revenue", "Escrow, document/exception and partner/API/ancillary fees after launch discounts."],
        ["Gross profit", "Revenue less modeled variable and shared COGS."],
        ["EBITDA before one-offs", "Gross profit less payroll and fixed opex; excludes startup/regulatory one-offs."],
        ["Operating result", "Gross profit less payroll, fixed opex and modeled one-offs; pre-tax and pre-financing."],
        ["Peak cash gap", "Largest cumulative monthly operating deficit before financing."],
        ["Operating reserve", "Spendable liquidity cushion included in the initial ask."],
        ["Restricted capital", "Regulatory capital held separately; not spendable runway."],
        ["Safeguarded client assets", "Gross customer escrow balance with an equal liability; 0% available to operations."],
    ]
    add_data_table(
        doc,
        ["Term", "Meaning in this report"],
        definitions,
        [1.65, 4.85],
        font_size=8.5,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_source_note(doc, "Source: Financial Model README, P&L 5yr, Launch Readiness, Cash Flow Statement and Client Funds Memo.")

    doc.add_paragraph("Material limitations", style="Heading 2")
    limitations = [
        ("1. No operating history.", "Deal volume, conversion, churn, CAC, review workload and holding period are assumptions or planning outputs, not observed performance."),
        ("2. No investment-grade valuation.", "The model is a five-year pre-tax, pre-financing operating plan. It does not model corporation tax, interest, depreciation/amortisation, working-capital accounting beyond the cash plan, or an enterprise valuation."),
        ("3. Year-4 and Year-5 are ambitious.", "The base case depends on partner distribution, self-serve onboarding, digital documents and an AI-enabled staffing model."),
        ("4. Regulatory figures are placeholders.", "The VARA application, supervision and restricted-capital assumptions require confirmation from counsel, the partner and regulators."),
        ("5. Client-fund timing is provisional.", "Average safeguarded balances use a 14-day funded-to-release period. A 30-, 60- or 90-day cycle materially increases safeguarding scale but still creates no operating cash."),
        ("6. Pricing is not willingness-to-pay evidence.", "The price book is informed by deal-value and competitor research; only a paid pilot can validate conversion and retention."),
    ]
    for lead, body in limitations:
        add_para(doc, f"{lead} {body}", bold_lead=lead)

    doc.add_paragraph("Financial conclusion", style="Heading 2")
    add_para(doc, "Prefunding avoids balance-sheet credit exposure. Tiered pricing covers the higher cost of manual workflows, while digital adoption allows the take rate to decline without removing transaction contribution. Client funds remain segregated from company cash.")
    add_para(doc, "The base case requires £3.063m, records first revenue after Month 12 and depends on a large Year-4 increase in distribution. Capital should first fund the regulated pilot and production controls. Further funding should follow observed price acceptance, workload, cycle time, repeat use and partner-led distribution.")
    add_callout(doc, "Financing recommendation", "Commit capital to the regulated pilot and production controls first. Release scale capital only after paid transactions validate the main commercial assumptions.", fill=LIGHT_TEAL, accent=TEAL)

    reference_section = doc.add_section(WD_SECTION.NEW_PAGE)
    reference_section.different_first_page_header_footer = False
    reference_section.header.is_linked_to_previous = True
    reference_section.footer.is_linked_to_previous = True

    # References and source map
    add_section_heading(doc, "References and model source map")
    references = [
        ["[1]", "Team Transakt, Blockmediary Financial Model, 13 August 2026."],
        ["[2]", "Blockmediary, Defensible Average Deal Values for Documentary Escrow, 11 August 2026, especially pp.2-6 and 24-25."],
        ["[3]", "Blockmediary, DIFC Partner-Led Pilot and VARA Scale Readiness Report, 12 August 2026, especially pp.1-3 and 11."],
        ["[4]", "Blockmediary, Competitor Analysis, August 2026, especially pp.2-3 and 12-13."],
        ["[5]", "Blockmediary, Legal and Compliance Risk Assessment, August 2026, especially pp.2-4 and 22-24."],
        ["[6]", "Blockmediary, Trade Escrow Agreement, August 2026, especially pp.2-4 and 18-19."],
        ["[7]", "Blockmediary repository README and Business Requirements Document, baselined 9 August 2026."],
    ]
    add_data_table(
        doc,
        ["Ref.", "Source"],
        references,
        [0.55, 5.95],
        font_size=7.9,
        first_col_bold=True,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        cell_margins_dxa={"top": 35, "bottom": 35, "start": 120, "end": 80},
    )

    doc.add_paragraph("Primary workbook locations", style="Heading 2")
    source_rows = [
        ["Executive headlines and use of funds", "Dashboard"],
        ["Launch timing, regulatory costs and initial funding", "Launch Readiness!A5:J23"],
        ["Tier price book, revenue, cost and contribution", "Tier Economics!B3:U35"],
        ["Five-year operating case and customer bridge", "P&L 5yr!C7:G91"],
        ["Monthly cash flow and break-even", "Cash Flow Statement!B9:BI58"],
        ["Low / Base / High cases", "Scenario Engine!A1:R64"],
        ["Safeguarded client assets", "Client Funds Memo!A5:H31"],
        ["Mechanical validation", "Model Checks!A1:G48"],
        ["Model inputs and source notes", "Assumptions; Driver Tables; Risk Register"],
    ]
    add_data_table(
        doc,
        ["Topic", "Workbook location"],
        source_rows,
        [3.5, 3.0],
        font_size=7.9,
        first_col_bold=False,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        cell_margins_dxa={"top": 35, "bottom": 35, "start": 120, "end": 80},
    )
    add_source_note(doc, "Financial figures are drawn from the Blockmediary Financial Model dated 13 August 2026 and rounded for presentation. The workbook remains the numerical source of truth.")

    return doc


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    save_charts(data)
    doc = build_report(data)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
