from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


path = Path(r"C:\Users\cwbec\BlockMed\docs\Blockmediary_Deal_Value_Research_Report_Submission_Edition.docx")
doc = Document(path)

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)} sections={len(doc.sections)}")
for index, section in enumerate(doc.sections, 1):
    print(
        f"section{index}: width={section.page_width.inches:.2f}in height={section.page_height.inches:.2f}in "
        f"top={section.top_margin.inches:.2f} bottom={section.bottom_margin.inches:.2f} "
        f"left={section.left_margin.inches:.2f} right={section.right_margin.inches:.2f} "
        f"orientation={section.orientation}"
    )

hidden_paragraphs = 0
page_break_before = 0
keep_next = 0
keep_lines = 0
for paragraph in doc.paragraphs:
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None:
        page_break_before += 1
    if ppr is not None and ppr.find(qn("w:keepNext")) is not None:
        keep_next += 1
    if ppr is not None and ppr.find(qn("w:keepLines")) is not None:
        keep_lines += 1
    runs = paragraph.runs
    if runs and all(
        run._r.rPr is not None and run._r.rPr.find(qn("w:vanish")) is not None
        for run in runs
    ):
        hidden_paragraphs += 1

with ZipFile(path) as archive:
    bad = archive.testzip()
    settings = archive.read("word/settings.xml").decode("utf-8")
    document_xml = archive.read("word/document.xml").decode("utf-8")

print(
    f"zip_bad_member={bad!r} hidden_paragraphs={hidden_paragraphs} "
    f"page_break_before={page_break_before} keep_next={keep_next} keep_lines={keep_lines}"
)
print(f"document_xml_chars={len(document_xml)} compat15={'compatibilityMode' in settings and 'w:val=\"15\"' in settings}")
print(f"last_paragraph={doc.paragraphs[-1].text[:120]!r}")
